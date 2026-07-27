"""DOCX generator (T-34) — python-docx.

Word's built-in paragraph styles carry the document structure (real headings,
real bullet lists — so Word's navigation pane and outline view work), with the
app palette applied on top: accent headings, muted timestamps, struck-through
completed tasks. Checkboxes are the Unicode ballot-box glyphs; a form control
would need Word's content-control XML for something a reader never toggles.
"""

from __future__ import annotations

import io
from typing import TYPE_CHECKING, assert_never

from docx import Document
from docx.shared import Pt, RGBColor

from app.services.export import blocks, palette
from app.services.export.blocks import clock, task_suffix

if TYPE_CHECKING:
    from docx.document import Document as DocumentObject
    from docx.text.paragraph import Paragraph as DocxParagraph

    from app.services.export.blocks import Block, ExportDocument

#: U+2610 ballot box / U+2611 ballot box with check.
_BOX_OPEN = "☐"
_BOX_CHECKED = "☑"


def _rgb(token: str) -> RGBColor:
    return RGBColor.from_string(token.lstrip("#"))


def render_docx(document: ExportDocument) -> bytes:
    doc = Document()

    title = doc.add_heading(document.title, level=0)
    _tint(title, _rgb(palette.INK))

    for label, value in document.metadata:
        paragraph = doc.add_paragraph()
        paragraph.add_run(f"{label}: ").bold = True
        paragraph.add_run(value)

    for block in document.blocks:
        _add_block(doc, block)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _tint(paragraph: DocxParagraph, color: RGBColor) -> None:
    """Word's default heading blue is not our palette; re-ink every run."""
    for run in paragraph.runs:
        run.font.color.rgb = color


def _add_block(doc: DocumentObject, block: Block) -> None:
    match block:
        case blocks.Heading(text):
            _tint(doc.add_heading(text, level=1), _rgb(palette.ACCENT))
        case blocks.Subheading(text):
            _tint(doc.add_heading(text, level=2), _rgb(palette.SECONDARY))
        case blocks.Paragraph(text):
            doc.add_paragraph(text)
        case blocks.Bullets(items):
            for item in items:
                doc.add_paragraph(item, style="List Bullet")
        case blocks.Outline(entries):
            for entry in entries:
                paragraph = doc.add_paragraph(style="List Bullet")
                stamp = paragraph.add_run(f"[{clock(entry.start_ms)}] ")
                stamp.font.color.rgb = _rgb(palette.MUTED)
                paragraph.add_run(entry.title)
        case blocks.Checklist(tasks):
            for task in tasks:
                _add_task(doc, task)
        case blocks.Transcript(turns):
            for turn in turns:
                paragraph = doc.add_paragraph()
                paragraph.add_run(turn.speaker).bold = True
                stamp = paragraph.add_run(f" [{clock(turn.start_ms)}] ")
                stamp.font.color.rgb = _rgb(palette.MUTED)
                paragraph.add_run(turn.text)
        case _:
            assert_never(block)


def _add_task(doc: DocumentObject, task: blocks.Task) -> None:
    paragraph = doc.add_paragraph()
    paragraph.add_run(f"{_BOX_CHECKED if task.done else _BOX_OPEN} ")
    body = paragraph.add_run(task.text)
    if task.done:
        body.font.strike = True
        body.font.color.rgb = _rgb(palette.MUTED)
    suffix = task_suffix(task)
    if suffix:
        meta = paragraph.add_run(suffix)
        meta.font.color.rgb = _rgb(palette.MUTED)
        meta.font.size = Pt(9)
